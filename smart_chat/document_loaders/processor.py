import os
import base64
import hashlib
import requests
import logging
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from smart_chat.core.config import CONFIG
from smart_chat.vectorstores.chroma import VectorStoreManager

logger = logging.getLogger(__name__)

# 确保临时目录存在
def ensure_temp_dir(temp_dir: str) -> str:
    """确保临时目录存在"""
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    return temp_dir


def get_poppler_path() -> str | None:
    """获取 Poppler bin 目录路径，用于 Windows 下 pdf2image 渲染 PDF。"""
    env_path = os.environ.get("POPPLER_PATH")
    if env_path and os.path.exists(env_path):
        return env_path

    candidates = [
        os.path.join(os.getcwd(), "tools", "poppler", "bin"),
        os.path.join(os.getcwd(), "tools", "poppler-23.11.0", "Library", "bin"),
    ]
    for p in candidates:
        if os.path.exists(p) and os.path.exists(os.path.join(p, "pdftoppm.exe")):
            return p
    return None

class DocumentProcessor:
    """文档处理类，支持 Word/PDF/图片等多种格式。"""
    
    def __init__(self):
        """初始化文档处理器。"""
        doc_config = CONFIG.get("document_processing", {})
        self.temp_dir = ensure_temp_dir(doc_config.get("temp_dir", "./temp"))
        self.vision_model = doc_config.get("vision_model", "granite3.2-vision:latest")
        self.ollama_api = doc_config.get("api_endpoint", "http://localhost:1314/api/generate")
        self.cache = {}
        self.cache_size = doc_config.get("cache_size", 10)
        self.vector_store = VectorStoreManager()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100
        )
        
    def _index_text_to_vectorstore(self, file_path: str, file_hash: str, text: str) -> bool:
        """将文本切分后写入向量库。
        
        Args:
            file_path: 文件路径
            file_hash: 文件 MD5 哈希
            text: 文档文本内容
            
        Returns:
            是否成功索引
        """
        if not text or text.startswith("Error"):
            return False

        try:
            file_name = os.path.basename(file_path)
            doc = Document(
                page_content=text,
                metadata={"source": file_name, "file_hash": file_hash},
            )
            splits = self.text_splitter.split_documents([doc])
            if not self.vector_store.ensure_ready():
                error_detail = self.vector_store.get_last_error() or "未知错误"
                logger.error(f"向量库初始化失败: {error_detail}")
                return False
            ok = self.vector_store.add_documents(splits)
            if ok:
                logger.info(f"已将文档 {file_name} 索引到向量库 ({len(splits)} 个片段)")
            return ok
        except Exception as e:
            logger.error(f"向量化文档失败: {e}", exc_info=True)
            return False

    def process_document(self, file_path: str, max_pages=10, index_to_kb: bool = True) -> str:
        """处理文档并提取文本。
        
        Args:
            file_path: 文件路径
            max_pages: PDF 最大处理页数
            index_to_kb: 是否索引到向量库
            
        Returns:
            提取的文本内容
        """
        if not file_path or not os.path.exists(file_path):
            return ""
            
        # 计算文件哈希值用于缓存
        with open(file_path, "rb") as f:
            file_content = f.read()
            file_hash = hashlib.md5(file_content).hexdigest()
        
        # 检查缓存中是否已有结果
        if file_hash in self.cache:
            cached = self.cache[file_hash]
            if isinstance(cached, dict):
                text_content = cached.get("text", "")
                indexed = bool(cached.get("indexed", False))
            else:
                text_content = cached
                indexed = False

            if index_to_kb and (not indexed):
                ok = self._index_text_to_vectorstore(file_path=file_path, file_hash=file_hash, text=text_content)
                if isinstance(cached, dict):
                    cached["indexed"] = ok
                else:
                    self.cache[file_hash] = {"text": text_content, "indexed": ok}

            return text_content
            
        # 没有缓存，处理文档
        text_content = self._process_document(file_hash, file_path, max_pages)
        
        indexed = False
        if index_to_kb:
            indexed = self._index_text_to_vectorstore(
                file_path=file_path, file_hash=file_hash, text=text_content
            )
            if (not indexed) and text_content and (not text_content.startswith("Error")):
                file_name = os.path.basename(file_path)
                return f"Error: 向量库未初始化或写入失败 ({file_name})"
        
        # 更新缓存
        self.cache[file_hash] = {"text": text_content, "indexed": indexed}
        
        # 如果缓存太大，移除最早添加的项
        if len(self.cache) > self.cache_size:
            oldest_key = next(iter(self.cache))
            del self.cache[oldest_key]
            
        return text_content
    
    def _process_document(self, file_hash, file_path, max_pages):
        """处理文档的内部函数"""
        # 获取文件扩展名（小写）
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 根据文件类型处理
        if file_ext in ['.docx', '.doc']:
            # Word文档处理
            return self._process_word_document(file_path)
            
        elif file_ext == '.pdf':
            # PDF处理
            return self._process_pdf_document(file_path, max_pages)
            
        elif file_ext in ['.txt', '.md', '.csv']:
            # 文本文件直接读取
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # 如果UTF-8解码失败，尝试其他编码
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        return f.read()
                except:
                    # 如果还是失败，将文件作为图像处理
                    return self._process_images([file_path])
        else:
            # 图片和其他文件作为图像处理
            return self._process_images([file_path])
    
    def _process_word_document(self, doc_path):
        """处理Word文档并提取文本"""
        try:
            # 尝试导入docx库
            try:
                import docx
                document = docx.Document(doc_path)
                # 提取文本
                text = "\n\n".join([para.text for para in document.paragraphs if para.text.strip()])
                if text:
                    return text
            except ImportError:
                logger.warning("docx module not found, fallback to image processing")
                return self._process_images([doc_path])
            except Exception as e:
                logger.warning(f"Failed to extract text with python-docx: {e}")
            
            # 如果文本提取失败，使用图像处理
            return self._process_images([doc_path])
                
        except Exception as e:
            error_msg = f"Error processing Word document: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _process_pdf_document(self, pdf_path, max_pages):
        """处理PDF文档并提取文本"""
        try:
            # 尝试使用pdf2image提取图像
            try:
                from pdf2image import convert_from_path
                
                # 转换PDF为图片
                poppler_path = get_poppler_path()
                try:
                    images = convert_from_path(
                        pdf_path,
                        first_page=1,
                        last_page=max_pages,
                        poppler_path=poppler_path,
                    )
                except Exception as e:
                    logger.warning(f"pdf2image 转换失败，将回退到文本提取: {e}")
                    raise ImportError("pdf2image failed") from e
                
                # 成功转换
                image_paths = []
                
                # 保存转换后的图片
                for i, img in enumerate(images):
                    img_path = os.path.join(self.temp_dir, f"page_{i}_{os.path.basename(pdf_path)}.jpg")
                    img.save(img_path, "JPEG")
                    image_paths.append(img_path)
                
                # 处理所有图片
                return self._process_images(image_paths)
                
            except ImportError:
                # 如果pdf2image不可用，尝试使用PyPDF2
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(pdf_path)
                    text = ""
                    for i, page in enumerate(reader.pages):
                        if i >= max_pages:
                            break
                        text += page.extract_text() + "\n\n"
                    
                    if text.strip():
                        return text
                    else:
                        return (
                            "Error: PDF 未提取到可用文本。若为扫描版/图片型 PDF，"
                            "请安装 Poppler 并配置 POPPLER_PATH 以启用 pdf2image 渲染。"
                        )
                        
                except ImportError:
                    # 如果两个库都不可用，使用图像处理
                    logger.warning("No PDF library available")
                    return (
                        "Error: 缺少 PDF 解析依赖。文本型 PDF 请安装 PyPDF2，扫描版/图片型 PDF 需安装 Poppler+pdf2image。"
                    )
                    
        except Exception as e:
            error_msg = f"Error processing PDF document: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _process_images(self, image_paths):
        """处理图片列表并提取文本内容"""
        parsed_text = []
        
        for i, img_path in enumerate(image_paths):
            try:
                # 转换图片为base64
                with open(img_path, "rb") as f:
                    img_base64 = base64.b64encode(f.read()).decode()
                
                # 构建提示词
                prompt = "Extract all text content from this document, maintaining the original format"
                
                # 调用模型解析
                response = requests.post(
                    self.ollama_api,
                    json={
                        "model": self.vision_model,
                        "prompt": prompt,
                        "images": [img_base64],
                        "stream": False
                    },
                    timeout=60
                )
                
                # 检查响应状态
                if response.status_code == 200:
                    result = response.json()
                    parsed_text.append(result.get("response", ""))
                else:
                    error_msg = f"API request failed: HTTP {response.status_code}"
                    parsed_text.append(error_msg)
                    
            except Exception as e:
                error_msg = f"Error processing image: {str(e)}"
                parsed_text.append(error_msg)
                
        return "\n\n".join(parsed_text)

# 创建单例实例
document_processor = DocumentProcessor()
